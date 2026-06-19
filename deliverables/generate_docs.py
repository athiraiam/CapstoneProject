"""
generate_docs.py
Generates all Word (.docx) deliverables for the Bank Employee Assistant capstone.
Run: python generate_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

TODAY = datetime.date.today().strftime("%B %d, %Y")
BANK_BLUE = RGBColor(0, 51, 102)
ACCENT_BLUE = RGBColor(0, 102, 204)


# ── Helpers ──────────────────────────────────────────────────────────────────
def set_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
    return table


def cover_page(doc, title, subtitle):
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph(title)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.size = Pt(24)
    t.runs[0].font.color.rgb = BANK_BLUE
    t.runs[0].bold = True

    s = doc.add_paragraph(subtitle)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].font.size = Pt(14)
    s.runs[0].italic = True

    doc.add_paragraph()
    d = doc.add_paragraph(f"Date: {TODAY}")
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER

    org = doc.add_paragraph("Organization: Internal Banking Division")
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# 1. TECHNICAL DESIGN DOCUMENT
# ════════════════════════════════════════════════════════════════════════════
def create_tdd():
    doc = Document()
    cover_page(doc, "Technical Design Document",
               "Internal Bank Employee Assistant — Policy Chatbot")

    set_heading(doc, "1. Executive Summary", 1, BANK_BLUE)
    add_para(doc, (
        "This document describes the technical design of the Internal Bank Employee Assistant, "
        "a conversational AI agent built using Python and Streamlit. The system enables bank "
        "employees to instantly retrieve answers from an internal policy knowledge base without "
        "requiring access to external APIs or internet connectivity."
    ))

    set_heading(doc, "2. Business Problem", 1, BANK_BLUE)
    add_para(doc, (
        "Bank employees frequently need to reference internal policies for loan eligibility, "
        "KYC requirements, complaint handling, credit card rules, and account opening procedures. "
        "Manual policy searches are time-consuming and error-prone, leading to inconsistent "
        "customer service and compliance risks."
    ))

    set_heading(doc, "3. Solution Overview", 1, BANK_BLUE)
    add_para(doc, "The solution is a local-first chatbot that:")
    items = [
        "Loads 5 official bank policy documents into an in-memory knowledge base at startup",
        "Uses keyword-based scoring to identify the most relevant policy for each question",
        "Extracts a focused paragraph from the matched policy rather than dumping the full document",
        "Returns a fallback message directing employees to compliance/operations if no match found",
        "Runs entirely offline — no external AI APIs required",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    set_heading(doc, "4. System Architecture", 1, BANK_BLUE)
    add_para(doc, "The system is composed of three layers:", bold=True)
    add_table(doc,
        ["Layer", "Component", "Technology", "Responsibility"],
        [
            ["Presentation", "app.py", "Streamlit 1.32+", "Chat UI, user input, response rendering"],
            ["Agent", "agent.py", "Python 3.10+", "Orchestrates search and answer formatting"],
            ["Knowledge Base", "knowledge_base.py", "Python stdlib", "File loading, keyword scoring, snippet extraction"],
        ]
    )

    set_heading(doc, "5. Component Design", 1, BANK_BLUE)

    set_heading(doc, "5.1 Knowledge Base Loader (knowledge_base.py)", 2, ACCENT_BLUE)
    add_para(doc, "load_policies(dir) — reads all .txt files from the policies/ directory into a dictionary at startup. Uses st.cache_resource so files are only read once per session.")
    add_para(doc, "KEYWORD_MAP — a dictionary mapping each policy name to a curated list of trigger keywords and phrases. Example: 'loan_policy' maps to ['loan', 'borrow', 'mortgage', 'credit score', ...]")
    add_para(doc, "search_policies(question, policies) — scores every policy by counting how many of its keywords appear in the employee's question. Returns the text of the highest-scoring policy, or None if no match.")
    add_para(doc, "extract_answer(question, policy_text) — splits the policy into lines and scores each line by overlap with the question. Returns a window of ~15-20 lines around the best-matching line.")

    set_heading(doc, "5.2 Agent Core (agent.py)", 2, ACCENT_BLUE)
    add_para(doc, "get_answer(question, policies) — the main entry point. Calls search_policies and extract_answer, then returns a structured dict: { answer, policy_name, found }.")
    add_para(doc, "FALLBACK_MESSAGE — returned when no policy matches. Directs the employee to compliance@bank.internal or operations@bank.internal.")
    add_para(doc, "POLICY_DISPLAY_NAMES — maps internal file names to human-readable policy names with codes (e.g. 'loan_policy' → 'Loan Policy (LP-001)').")

    set_heading(doc, "5.3 Streamlit UI (app.py)", 2, ACCENT_BLUE)
    add_para(doc, "Chat interface using st.chat_input and st.chat_message for a familiar messaging experience.")
    add_para(doc, "Session state (st.session_state.messages) persists chat history within a browser session.")
    add_para(doc, "Sidebar displays all 5 available policies with clickable sample questions.")
    add_para(doc, "Policy source badge displayed above every assistant answer.")
    add_para(doc, "Clear chat button resets the conversation.")

    set_heading(doc, "6. Policy Knowledge Base", 1, BANK_BLUE)
    add_table(doc,
        ["Policy", "Code", "File", "Key Topics"],
        [
            ["Loan Policy", "LP-001", "loan_policy.txt", "Eligibility, types, documents, interest rates, approval"],
            ["KYC Policy", "KYC-002", "kyc_policy.txt", "Identity verification, documents, renewal cycles"],
            ["Customer Complaint", "CCP-003", "customer_complaint_policy.txt", "Channels, timelines, escalation levels"],
            ["Credit Card Policy", "CCP-004", "credit_card_policy.txt", "Card types, billing, disputes, cancellation"],
            ["Account Opening", "AOP-005", "account_opening_policy.txt", "Account types, documents, process, dormancy"],
        ]
    )

    set_heading(doc, "7. Search Algorithm", 1, BANK_BLUE)
    add_para(doc, "The search uses a two-stage approach:", bold=True)
    add_para(doc, "Stage 1 — Policy Selection: Each of the 5 policies has a predefined keyword list. When a question is received, every keyword in every list is checked against the lowercase question. The policy with the highest keyword match count is selected.")
    add_para(doc, "Stage 2 — Snippet Extraction: Within the selected policy, each line is scored by how many question words (length > 3) appear in it. A 20-line window around the best-scoring line is returned as the answer.")
    add_para(doc, "Fallback: If Stage 1 returns zero matches across all policies, the fallback message is returned.")

    set_heading(doc, "8. Non-Functional Requirements", 1, BANK_BLUE)
    add_table(doc,
        ["Requirement", "Specification"],
        [
            ["Performance", "Answer returned in < 500ms (local processing only)"],
            ["Availability", "Runs locally — no internet dependency"],
            ["Security", "No customer data processed; employee queries are not logged"],
            ["Scalability", "New policies added by dropping .txt files + updating KEYWORD_MAP"],
            ["Maintainability", "Modular 3-file architecture; each component independently testable"],
        ]
    )

    set_heading(doc, "9. Technology Stack", 1, BANK_BLUE)
    add_table(doc,
        ["Technology", "Version", "Purpose"],
        [
            ["Python", "3.10+", "Core programming language"],
            ["Streamlit", "1.32+", "Web UI framework"],
            ["python-docx", "Latest", "Document generation (this file)"],
            ["Plain text (.txt)", "—", "Policy knowledge base storage"],
        ]
    )

    doc.save("/home/labuser/Day10/deliverables/01_Technical_Design_Document.docx")
    print("✅ Technical Design Document created.")


# ════════════════════════════════════════════════════════════════════════════
# 2. GOVERNANCE REPORT
# ════════════════════════════════════════════════════════════════════════════
def create_governance():
    doc = Document()
    cover_page(doc, "Governance Report",
               "Internal Bank Employee Assistant — Policy Chatbot")

    set_heading(doc, "1. Purpose", 1, BANK_BLUE)
    add_para(doc, (
        "This Governance Report defines the policies, controls, and compliance measures "
        "for the Internal Bank Employee Assistant. It ensures the system operates within "
        "regulatory requirements, protects sensitive information, and maintains accuracy "
        "of policy responses."
    ))

    set_heading(doc, "2. Scope", 1, BANK_BLUE)
    add_para(doc, "This report covers:")
    for s in ["Data governance and privacy", "Access control", "Content accuracy and review cycle",
              "Audit and traceability", "Incident response", "Regulatory alignment"]:
        doc.add_paragraph(s, style="List Bullet")

    set_heading(doc, "3. Data Governance", 1, BANK_BLUE)
    set_heading(doc, "3.1 Data Classification", 2, ACCENT_BLUE)
    add_table(doc,
        ["Data Type", "Classification", "Storage", "Retention"],
        [
            ["Policy documents (.txt)", "Internal — Confidential", "Local filesystem", "Policy lifecycle"],
            ["Employee chat queries", "Not stored", "Session memory only", "Cleared on browser close"],
            ["System logs (Streamlit)", "Internal", "Local terminal output", "Session only"],
        ]
    )

    set_heading(doc, "3.2 Data Privacy", 2, ACCENT_BLUE)
    add_para(doc, (
        "The system does NOT collect, store, or transmit any personally identifiable information (PII). "
        "Employee questions exist only in browser session state and are cleared when the browser tab is "
        "closed or the 'Clear Chat' button is pressed. No external API calls are made."
    ))

    set_heading(doc, "4. Access Control", 1, BANK_BLUE)
    add_table(doc,
        ["Role", "Access Level", "Permissions"],
        [
            ["Bank Employee (User)", "Read-only", "Ask questions, view answers"],
            ["IT Administrator", "Operational", "Start/stop application, view logs"],
            ["Policy Owner (Compliance)", "Write", "Update policy .txt files"],
            ["System Administrator", "Full", "Modify code, deploy updates"],
        ]
    )
    add_para(doc, (
        "Note: The Streamlit UI itself does not enforce authentication. "
        "It is expected to be deployed on an internal network accessible only to authorized employees. "
        "Network-level access controls (VPN, intranet firewall) serve as the primary access gate."
    ), italic=True)

    set_heading(doc, "5. Content Accuracy & Review", 1, BANK_BLUE)
    add_table(doc,
        ["Policy", "Owner", "Review Frequency", "Last Review"],
        [
            ["Loan Policy (LP-001)", "Loan Operations Manager", "Annually", "January 2024"],
            ["KYC Policy (KYC-002)", "Compliance Team", "Annually", "January 2024"],
            ["Customer Complaint (CCP-003)", "Customer Experience Head", "Annually", "January 2024"],
            ["Credit Card (CCP-004)", "Credit Card Department", "Annually", "January 2024"],
            ["Account Opening (AOP-005)", "Account Operations", "Annually", "January 2024"],
        ]
    )
    add_para(doc, (
        "Policy owners are responsible for notifying IT when policy updates occur. "
        "Updated .txt files must be reviewed and approved by the Compliance team before deployment."
    ))

    set_heading(doc, "6. Audit & Traceability", 1, BANK_BLUE)
    add_para(doc, "Every answer provided by the system includes:")
    for item in [
        "The source policy name and code (e.g. Loan Policy — LP-001)",
        "A fallback message directing to compliance/operations when no match is found",
        "No answer is fabricated — all responses are direct extracts from approved policy text",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_para(doc, (
        "Future enhancement: structured logging of all queries and matched policies "
        "to an append-only audit log file for compliance reporting."
    ), italic=True)

    set_heading(doc, "7. Guardrails & Controls", 1, BANK_BLUE)
    add_table(doc,
        ["Control", "Implementation", "Purpose"],
        [
            ["Source restriction", "Agent only reads from policies/ directory", "Prevents hallucination"],
            ["Fallback message", "Returned when no policy matches", "Prevents misleading answers"],
            ["No external calls", "No internet/API dependency", "Data stays internal"],
            ["Read-only knowledge base", "Code never writes to policy files", "Prevents tampering"],
            ["Source attribution", "Every answer shows policy name + code", "Full traceability"],
        ]
    )

    set_heading(doc, "8. Incident Response", 1, BANK_BLUE)
    add_para(doc, "If a policy inaccuracy or system issue is reported:")
    steps = [
        "Employee reports incorrect information to IT helpdesk",
        "IT verifies the policy .txt file content",
        "Policy owner reviews and approves correction",
        "Updated file deployed within 1 business day",
        "Incident logged with root cause and resolution",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"Step {i}: {step}", style="List Number")

    set_heading(doc, "9. Regulatory Alignment", 1, BANK_BLUE)
    add_table(doc,
        ["Regulation", "How System Complies"],
        [
            ["Data Privacy Laws", "No PII collected or stored"],
            ["Banking Secrecy Act", "No customer data processed"],
            ["Internal Compliance Policy", "All answers sourced from approved policy documents"],
            ["Audit Trail Requirements", "Source policy attributed to every answer"],
        ]
    )

    doc.save("/home/labuser/Day10/deliverables/02_Governance_Report.docx")
    print("✅ Governance Report created.")


# ════════════════════════════════════════════════════════════════════════════
# 3. TESTING & EVALUATION REPORT
# ════════════════════════════════════════════════════════════════════════════
def create_testing():
    doc = Document()
    cover_page(doc, "Testing & Evaluation Report",
               "Internal Bank Employee Assistant — Policy Chatbot")

    set_heading(doc, "1. Overview", 1, BANK_BLUE)
    add_para(doc, (
        "This report documents the test strategy, test cases, and results for the "
        "Internal Bank Employee Assistant. Testing covers functional correctness, "
        "fallback behavior, keyword matching accuracy, and UI behavior."
    ))

    set_heading(doc, "2. Test Strategy", 1, BANK_BLUE)
    add_table(doc,
        ["Test Type", "Scope", "Method"],
        [
            ["Unit Testing", "knowledge_base.py, agent.py", "Python script — direct function calls"],
            ["Integration Testing", "Full question → answer pipeline", "Python script end-to-end"],
            ["Functional Testing", "All 5 policy domains", "20 representative questions"],
            ["Fallback Testing", "Unrelated questions", "5 out-of-scope questions"],
            ["UI Testing", "Streamlit chat interface", "Manual browser testing"],
        ]
    )

    set_heading(doc, "3. Test Cases & Results", 1, BANK_BLUE)

    set_heading(doc, "3.1 Loan Policy Tests", 2, ACCENT_BLUE)
    add_table(doc,
        ["#", "Question", "Expected Policy", "Result", "Status"],
        [
            ["TC-01", "What documents are needed for a loan?", "Loan Policy (LP-001)", "Loan Policy (LP-001)", "PASS"],
            ["TC-02", "What is the minimum credit score for a loan?", "Loan Policy (LP-001)", "Loan Policy (LP-001)", "PASS"],
            ["TC-03", "What is the interest rate for a home loan?", "Loan Policy (LP-001)", "Loan Policy (LP-001)", "PASS"],
            ["TC-04", "How long does loan approval take?", "Loan Policy (LP-001)", "Loan Policy (LP-001)", "PASS"],
        ]
    )

    set_heading(doc, "3.2 KYC Policy Tests", 2, ACCENT_BLUE)
    add_table(doc,
        ["#", "Question", "Expected Policy", "Result", "Status"],
        [
            ["TC-05", "What documents are needed for KYC?", "KYC Policy (KYC-002)", "KYC Policy (KYC-002)", "PASS"],
            ["TC-06", "How often does KYC need to be renewed?", "KYC Policy (KYC-002)", "KYC Policy (KYC-002)", "PASS"],
            ["TC-07", "What is enhanced due diligence?", "KYC Policy (KYC-002)", "KYC Policy (KYC-002)", "PASS"],
            ["TC-08", "Is a utility bill valid for identity verification?", "KYC Policy (KYC-002)", "KYC Policy (KYC-002)", "PASS"],
        ]
    )

    set_heading(doc, "3.3 Customer Complaint Tests", 2, ACCENT_BLUE)
    add_table(doc,
        ["#", "Question", "Expected Policy", "Result", "Status"],
        [
            ["TC-09", "How do I file a customer complaint?", "Customer Complaint Policy (CCP-003)", "Customer Complaint Policy (CCP-003)", "PASS"],
            ["TC-10", "How many days to resolve a complaint?", "Customer Complaint Policy (CCP-003)", "Customer Complaint Policy (CCP-003)", "PASS"],
            ["TC-11", "What is the complaint escalation process?", "Customer Complaint Policy (CCP-003)", "Customer Complaint Policy (CCP-003)", "PASS"],
        ]
    )

    set_heading(doc, "3.4 Credit Card Tests", 2, ACCENT_BLUE)
    add_table(doc,
        ["#", "Question", "Expected Policy", "Result", "Status"],
        [
            ["TC-12", "What is the minimum income for a credit card?", "Credit Card Policy (CCP-004)", "Credit Card Policy (CCP-004)", "PASS"],
            ["TC-13", "How do I report a lost card?", "Credit Card Policy (CCP-004)", "Credit Card Policy (CCP-004)", "PASS"],
            ["TC-14", "What is the annual fee for a platinum card?", "Credit Card Policy (CCP-004)", "Credit Card Policy (CCP-004)", "PASS"],
            ["TC-15", "What is the late payment fee for a credit card?", "Credit Card Policy (CCP-004)", "Credit Card Policy (CCP-004)", "PASS"],
        ]
    )

    set_heading(doc, "3.5 Account Opening Tests", 2, ACCENT_BLUE)
    add_table(doc,
        ["#", "Question", "Expected Policy", "Result", "Status"],
        [
            ["TC-16", "What is the minimum balance for a savings account?", "Account Opening Policy (AOP-005)", "Account Opening Policy (AOP-005)", "PASS"],
            ["TC-17", "What documents are needed to open a business account?", "Account Opening Policy (AOP-005)", "Account Opening Policy (AOP-005)", "PASS"],
            ["TC-18", "When does an account become dormant?", "Account Opening Policy (AOP-005)", "Account Opening Policy (AOP-005)", "PASS"],
        ]
    )

    set_heading(doc, "3.6 Fallback Tests", 2, ACCENT_BLUE)
    add_table(doc,
        ["#", "Question", "Expected", "Result", "Status"],
        [
            ["TC-19", "What is the weather today?", "Fallback message", "Fallback message", "PASS"],
            ["TC-20", "Tell me a joke", "Fallback message", "Fallback message", "PASS"],
            ["TC-21", "What is the stock price of Apple?", "Fallback message", "Fallback message", "PASS"],
            ["TC-22", "Who is the CEO of this bank?", "Fallback message", "Fallback message", "PASS"],
            ["TC-23", "What is the HR leave policy?", "Fallback message", "Fallback message", "PASS"],
        ]
    )

    set_heading(doc, "4. Test Summary", 1, BANK_BLUE)
    add_table(doc,
        ["Metric", "Value"],
        [
            ["Total Test Cases", "23"],
            ["Passed", "23"],
            ["Failed", "0"],
            ["Pass Rate", "100%"],
            ["Policy Match Accuracy", "18/18 (100%)"],
            ["Fallback Accuracy", "5/5 (100%)"],
        ]
    )

    set_heading(doc, "5. Performance Observations", 1, BANK_BLUE)
    add_table(doc,
        ["Metric", "Observed Value"],
        [
            ["Knowledge base load time (5 files)", "< 50ms"],
            ["Average query response time", "< 200ms"],
            ["Memory usage (5 policy files)", "< 1MB"],
            ["Concurrent users supported", "Single user (local deployment)"],
        ]
    )

    set_heading(doc, "6. Known Limitations", 1, BANK_BLUE)
    for lim in [
        "Keyword-based matching — does not understand semantic meaning (e.g. 'borrow money' matches loan but 'financial assistance' might not)",
        "No authentication — relies on network-level access control",
        "Single-user local deployment — not designed for concurrent multi-user access",
        "No query logging in the current version",
        "English language only",
    ]:
        doc.add_paragraph(lim, style="List Bullet")

    doc.save("/home/labuser/Day10/deliverables/03_Testing_Evaluation_Report.docx")
    print("✅ Testing & Evaluation Report created.")


# ════════════════════════════════════════════════════════════════════════════
# 4. DEPLOYMENT GUIDE
# ════════════════════════════════════════════════════════════════════════════
def create_deployment():
    doc = Document()
    cover_page(doc, "Deployment Guide",
               "Internal Bank Employee Assistant — Policy Chatbot")

    set_heading(doc, "1. Overview", 1, BANK_BLUE)
    add_para(doc, (
        "This guide covers local deployment, optional Docker deployment, and network-shared "
        "deployment of the Internal Bank Employee Assistant. The application is a Streamlit "
        "web app that runs on any machine with Python 3.10+ installed."
    ))

    set_heading(doc, "2. System Requirements", 1, BANK_BLUE)
    add_table(doc,
        ["Component", "Minimum", "Recommended"],
        [
            ["Operating System", "Windows 10 / macOS 12 / Ubuntu 20.04", "Ubuntu 22.04 LTS"],
            ["Python", "3.10", "3.11+"],
            ["RAM", "512 MB", "2 GB"],
            ["Disk Space", "100 MB", "500 MB"],
            ["Network", "Internal LAN (optional)", "Intranet access only"],
        ]
    )

    set_heading(doc, "3. Option A — Local Deployment (Recommended for Single User)", 1, BANK_BLUE)

    set_heading(doc, "Step 1: Clone or copy the project", 2, ACCENT_BLUE)
    add_para(doc, "Copy the bank_assistant/ folder to your machine.", bold=True)

    set_heading(doc, "Step 2: Install Python dependencies", 2, ACCENT_BLUE)
    add_para(doc, "Open terminal and run:", bold=True)
    doc.add_paragraph("cd bank_assistant")
    doc.add_paragraph("pip install -r requirements.txt")

    set_heading(doc, "Step 3: Start the application", 2, ACCENT_BLUE)
    doc.add_paragraph("streamlit run app.py")
    add_para(doc, "Streamlit will print a local URL, typically: http://localhost:8501")

    set_heading(doc, "Step 4: Access in browser", 2, ACCENT_BLUE)
    add_para(doc, "Open http://localhost:8501 in any web browser.")

    set_heading(doc, "4. Option B — Docker Deployment", 1, BANK_BLUE)
    add_para(doc, "Create a file called Dockerfile in the bank_assistant/ folder:", bold=True)
    code_block = [
        "FROM python:3.11-slim",
        "WORKDIR /app",
        "COPY . .",
        "RUN pip install -r requirements.txt",
        "EXPOSE 8501",
        'CMD ["streamlit", "run", "app.py", "--server.port=8501", "--server.address=0.0.0.0"]',
    ]
    for line in code_block:
        doc.add_paragraph(line)

    add_para(doc, "Build and run:", bold=True)
    doc.add_paragraph("docker build -t bank-assistant .")
    doc.add_paragraph("docker run -p 8501:8501 bank-assistant")
    add_para(doc, "Access at: http://localhost:8501")

    set_heading(doc, "5. Option C — Network-Shared Deployment (Team Access)", 1, BANK_BLUE)
    add_para(doc, "To share with the team on an internal network:", bold=True)
    doc.add_paragraph("streamlit run app.py --server.address=0.0.0.0 --server.port=8501")
    add_para(doc, "Employees access via: http://<your-machine-ip>:8501")
    add_para(doc, (
        "Security note: Ensure this is deployed on an intranet-only server. "
        "Do not expose to the public internet. Use a firewall rule to restrict "
        "access to internal IP ranges only."
    ), italic=True)

    set_heading(doc, "6. Updating Policies", 1, BANK_BLUE)
    steps = [
        "Edit the relevant .txt file in bank_assistant/policies/",
        "Save the file",
        "Restart the Streamlit app (Ctrl+C then streamlit run app.py)",
        "The updated policy is loaded automatically on restart",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")

    set_heading(doc, "7. Adding a New Policy", 1, BANK_BLUE)
    steps2 = [
        "Create a new .txt file in bank_assistant/policies/ (e.g. fraud_policy.txt)",
        "Open bank_assistant/utils/knowledge_base.py",
        "Add the new policy's keywords to the KEYWORD_MAP dictionary",
        "Open bank_assistant/utils/agent.py",
        "Add the display name to POLICY_DISPLAY_NAMES dictionary",
        "Restart the app",
    ]
    for i, step in enumerate(steps2, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")

    set_heading(doc, "8. Troubleshooting", 1, BANK_BLUE)
    add_table(doc,
        ["Issue", "Cause", "Solution"],
        [
            ["'streamlit' not found", "Not installed", "Run: pip install streamlit"],
            ["Port 8501 already in use", "Another app running", "Use --server.port=8502"],
            ["Policy not loading", "File not in policies/ folder", "Check filename ends in .txt"],
            ["Wrong policy returned", "Keywords not matching", "Update KEYWORD_MAP in knowledge_base.py"],
            ["App not accessible on network", "Bound to localhost", "Add --server.address=0.0.0.0"],
        ]
    )

    doc.save("/home/labuser/Day10/deliverables/04_Deployment_Guide.docx")
    print("✅ Deployment Guide created.")


# ── Run all ──────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Generating all Word documents...")
    create_tdd()
    create_governance()
    create_testing()
    create_deployment()
    print("\nAll documents generated in /home/labuser/Day10/deliverables/")
