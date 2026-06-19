"""
My_5Min_Presentation_Guide.py
-------------------------------
Generates a personal, human-touch reference doc for a 5-minute presentation.
Run: python My_5Min_Presentation_Guide.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

TODAY = datetime.date.today().strftime("%B %d, %Y")

# Colours
NAVY        = RGBColor(0,   51,  102)
BLUE        = RGBColor(0,   102, 204)
GREEN       = RGBColor(0,   128,  64)
ORANGE      = RGBColor(200,  90,   0)
GRAY        = RGBColor(90,   90,  90)
LIGHT_GRAY  = RGBColor(150, 150, 150)
BLACK       = RGBColor(0,     0,   0)


# ── Helpers ───────────────────────────────────────────────────────────────────
def section_title(doc, emoji, title, color=NAVY):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"{emoji}  {title}")
    run.bold           = True
    run.font.size      = Pt(14)
    run.font.color.rgb = color
    # Underline rule — a thin border on a blank paragraph
    border_para = doc.add_paragraph()
    border_para.paragraph_format.space_before = Pt(0)
    border_para.paragraph_format.space_after  = Pt(4)
    pPr = border_para._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0055b3')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def tip(doc, text):
    """A soft grey italic speaker tip."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f"💬  {text}")
    r.italic           = True
    r.font.size        = Pt(10)
    r.font.color.rgb   = LIGHT_GRAY
    return p


def body(doc, text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size      = Pt(size)
    r.font.color.rgb = BLACK
    return p


def bold_body(doc, label, text, label_color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label + "  ")
    r1.bold            = True
    r1.font.size       = Pt(11)
    r1.font.color.rgb  = label_color
    r2 = p.add_run(text)
    r2.font.size       = Pt(11)
    r2.font.color.rgb  = BLACK
    return p


def bullet(doc, items, color=BLACK):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after      = Pt(2)
        p.paragraph_format.left_indent      = Inches(0.2)
        r = p.add_run(item)
        r.font.size      = Pt(11)
        r.font.color.rgb = color


def mini_table(doc, headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold           = True
            run.font.color.rgb = NAVY
    for row_data in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = val
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()


def callout(doc, text, color=BLUE):
    """Highlighted callout box feel using a bold framed paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(f"★  {text}")
    r.bold           = True
    r.font.size      = Pt(11)
    r.font.color.rgb = color


# ════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ════════════════════════════════════════════════════════════════════════════
doc = Document()

# Global margins
for section in doc.sections:
    section.top_margin    = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)


# ── COVER ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()

cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover_title.add_run("NexaBank Policy Assistant")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY

cover_sub = doc.add_paragraph()
cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover_sub.add_run("My 5-Minute Presentation Cheat Sheet")
r.italic = True; r.font.size = Pt(14); r.font.color.rgb = BLUE

doc.add_paragraph()

cover_note = doc.add_paragraph()
cover_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover_note.add_run(
    "This is your personal quick-reference guide.\n"
    "Read it once, and you'll know exactly what to say for each section.\n"
    "Estimated talk time per section is in brackets."
)
r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GRAY

doc.add_paragraph()

# Time budget bar
time_box = doc.add_paragraph()
time_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = time_box.add_run("⏱  Total: 5 minutes  |  ~25–30 seconds per section")
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = GREEN

doc.add_page_break()


# ── QUICK MAP ─────────────────────────────────────────────────────────────────
section_title(doc, "🗺", "Quick Section Map  (your 5-minute roadmap)", NAVY)
body(doc, "Use this as your talk track. Each section has a one-liner you can say out loud.")
mini_table(doc,
    ["#", "Section", "Time", "One-Liner"],
    [
        ["1",  "Business Problem",          "~25s", "Why we built this"],
        ["2",  "Solution Overview",         "~25s", "What we built — the big picture"],
        ["3",  "Agent Architecture",        "~30s", "How the brain works"],
        ["4",  "Skills, Subagents & Hooks", "~30s", "The team inside the brain"],
        ["5",  "MCP & Plugin Integration",  "~20s", "How AI tools plug in"],
        ["6",  "Governance Framework",      "~20s", "Why it's safe and auditable"],
        ["7",  "Observability & Traceability","~20s","How we watch what it does"],
        ["8",  "Evaluation Results",        "~20s", "Did it pass all the tests?"],
        ["9",  "Load Testing Results",      "~20s", "Can it handle real traffic?"],
        ["10", "Deployment Architecture",   "~20s", "Where and how it runs"],
        ["11", "Screenshots of Results",    "~20s", "Show, don't tell"],
        ["12", "Business Impact",           "~20s", "Why it matters to the business"],
    ]
)


# ── 1. BUSINESS PROBLEM ───────────────────────────────────────────────────────
section_title(doc, "1️⃣", "Business Problem  [~25 seconds]", NAVY)
body(doc,
    "Every day, bank employees have to answer customer questions about loans, KYC, "
    "credit cards, account opening, and complaint procedures. To do that, they manually "
    "dig through 5 different policy documents — each one 50+ lines long. That takes "
    "5 to 10 minutes per lookup. Multiply that by dozens of employees and hundreds of "
    "queries a day, and you have a real operational problem."
)
bullet(doc, [
    "Inconsistent answers — different employees read the same policy differently",
    "No audit trail — nobody knows which queries went unanswered or were escalated",
    "New employees struggle the most — policy documents are not easy to navigate",
])
tip(doc, "Say it like this: 'Our employees were wasting 5-10 minutes every time "
         "a customer asked a simple policy question. We wanted to fix that.'")


# ── 2. SOLUTION OVERVIEW ──────────────────────────────────────────────────────
section_title(doc, "2️⃣", "Solution Overview  [~25 seconds]", NAVY)
body(doc,
    "We built an internal AI-powered chatbot — the NexaBank Policy Assistant. "
    "It runs entirely offline, inside the bank's network, with no external APIs or "
    "internet connection needed. Employees just type their question and get the "
    "answer in under a second, with the source policy shown right alongside."
)
callout(doc,
    "It's not just a chatbot — it's a multi-agent system with 7 specialised agents "
    "that classify, search, summarise, escalate, and audit every single query.")
bullet(doc, [
    "Covers all 5 official bank policies (Loan, KYC, Complaints, Credit Card, Account Opening)",
    "Runs on Streamlit — no installation needed for employees, just open a browser",
    "Fully offline — all data stays inside the bank",
])
tip(doc, "One sentence version: 'We replaced 5-10 minute manual searches with a "
         "one-second AI answer, built entirely in-house.'")


# ── 3. AGENT ARCHITECTURE ─────────────────────────────────────────────────────
section_title(doc, "3️⃣", "Agent Architecture  [~30 seconds]", NAVY)
body(doc,
    "The system isn't just one AI — it's an orchestrated pipeline of 7 specialised "
    "agents working together. Think of the OrchestratorAgent as the team manager. "
    "Every employee question goes through exactly 9 steps:"
)
mini_table(doc,
    ["Step", "What Happens"],
    [
        ["1", "pre_query_hook — validates and sanitises the input"],
        ["2", "QueryClassifierAgent — figures out what the employee actually wants"],
        ["3", "OrchestratorAgent — decides which agent to send it to"],
        ["4", "The right subagent runs and produces an answer"],
        ["5", "AuditLoggerAgent — silently logs everything in the background"],
        ["6", "compliance_alert_hook — fires if any risk keywords are detected"],
        ["7", "escalation_hook — fires if the question couldn't be answered"],
        ["8", "post_answer_hook — logs response time and outcome"],
        ["9", "Answer returned to the chat UI with source badge and agent chip"],
    ]
)
tip(doc, "Keep it simple: 'Every question flows through 9 steps automatically. "
         "The system classifies, routes, answers, logs, and checks compliance — "
         "all without the employee doing anything extra.'")


# ── 4. SKILLS, SUBAGENTS & HOOKS ─────────────────────────────────────────────
section_title(doc, "4️⃣", "Skills, Subagents & Hooks  [~30 seconds]", NAVY)

body(doc, "The 7 subagents each have a specific job:")
mini_table(doc,
    ["Subagent", "Job in Plain English"],
    [
        ["QueryClassifierAgent",  "Reads the question and decides: is this a search? a summary? an escalation?"],
        ["PolicySearchAgent",     "Finds the best-matching policy and returns the relevant paragraph"],
        ["MultiPolicyAgent",      "When a question spans two policies, it merges both answers"],
        ["SummaryAgent",          "Gives a clean 5-bullet overview of any policy on demand"],
        ["EscalationAgent",       "Routes to the right team and auto-drafts an escalation email"],
        ["FallbackAgent",         "Self-heals when nothing matches — gives a fallback and logs it"],
        ["AuditLoggerAgent",      "Silently logs every query, result, and compliance risk in the background"],
    ]
)

body(doc, "The 3 skills are reusable functions shared across agents:")
bullet(doc, [
    "search_skill.py — keyword scoring and snippet extraction",
    "format_skill.py — formatting, summaries, email drafting",
    "audit_skill.py — all logging and compliance risk detection",
])

body(doc, "The 6 hooks fire automatically at key moments in every query:")
bullet(doc, [
    "pre_query_hook → validates input before anything runs",
    "post_answer_hook → logs outcome and response time after every answer",
    "session_start_hook / session_end_hook → tracks session lifecycle",
    "escalation_hook → logs unanswered queries for compliance review",
    "compliance_alert_hook → fires when risk keywords are detected",
])
tip(doc, "Analogy: 'Subagents are like specialists in a hospital. The "
         "Orchestrator is the triage nurse who sends you to the right doctor.'")


# ── 5. MCP & PLUGIN INTEGRATION ───────────────────────────────────────────────
section_title(doc, "5️⃣", "MCP & Plugin Integration  [~20 seconds]", NAVY)
body(doc,
    "We also built an MCP (Model Context Protocol) server — mcp_server.py. "
    "This means the knowledge base isn't locked inside the Streamlit app. "
    "Any MCP-compatible AI client — like Claude Desktop or Claude Code — can "
    "call our bank policy tools directly, without opening the browser app."
)
mini_table(doc,
    ["MCP Tool", "What it does"],
    [
        ["search_policy",         "Search for any policy topic by keyword"],
        ["list_policies",         "List all 5 available policies"],
        ["get_policy_full",       "Get the complete text of any policy by code"],
        ["cross_reference",       "Find all policies that match a multi-domain question"],
        ["check_compliance_risk", "Screen a question for 14 compliance risk keywords"],
    ]
)
tip(doc, "Simple version: 'We future-proofed it. Any AI tool can now plug "
         "directly into our policy knowledge base via an open standard.'")


# ── 6. GOVERNANCE FRAMEWORK ───────────────────────────────────────────────────
section_title(doc, "6️⃣", "Governance Framework  [~20 seconds]", NAVY)
body(doc,
    "This was important to get right. The system never makes up an answer — "
    "every response is a verbatim extract from an approved policy document. "
    "Here's what the governance framework ensures:"
)
bullet(doc, [
    "No PII collected — employee questions exist only in browser memory, cleared on tab close",
    "No external APIs — all data stays inside the bank's network",
    "Source attribution on every answer — employees always know which policy was used",
    "Only AuditLoggerAgent writes to log files — all other agents are read-only",
    "14 compliance risk keywords are auto-detected and flagged in real time",
    "Policy files can only be updated by authorised policy owners",
])
callout(doc, "Zero hallucination risk — the system cannot answer from anything "
             "other than the approved policy text.", GREEN)
tip(doc, "For the audience: 'Every answer can be traced back to the exact "
         "paragraph in the exact approved document. Nothing is invented.'")


# ── 7. OBSERVABILITY & TRACEABILITY ───────────────────────────────────────────
section_title(doc, "7️⃣", "Observability & Traceability  [~20 seconds]", NAVY)
body(doc,
    "Every single query leaves a complete audit trail across 3 log files:"
)
mini_table(doc,
    ["Log File", "What's recorded"],
    [
        ["query_audit.log",       "Every query: timestamp, session, MATCHED/FALLBACK, policy used"],
        ["unresolved_queries.log","Every question that couldn't be answered — for gap analysis"],
        ["session_events.log",    "Full lifecycle: session start/end, escalations, compliance alerts"],
    ]
)
body(doc,
    "In the UI, each answer shows: which agent answered it, what intent was detected, "
    "and which policy was the source. Compliance risk queries trigger a visible yellow "
    "alert banner on screen."
)
tip(doc, "Key point: 'If a compliance officer asks what happened with any query "
         "at any time, we can pull the exact log entry in seconds.'")


# ── 8. EVALUATION RESULTS ────────────────────────────────────────────────────
section_title(doc, "8️⃣", "Evaluation Results  [~20 seconds]", NAVY)
body(doc,
    "We ran a comprehensive test suite — 49 test cases across every agent, "
    "every skill, every hook, and the full end-to-end pipeline."
)
callout(doc, "49 tests.  49 passed.  0 failed.  100% pass rate.", GREEN)
mini_table(doc,
    ["Test Category", "Tests", "Result"],
    [
        ["QueryClassifierAgent",             "6",  "6/6 ✅"],
        ["PolicySearchAgent (all 5 domains)","5",  "5/5 ✅"],
        ["MultiPolicyAgent",                 "4",  "4/4 ✅"],
        ["SummaryAgent",                     "4",  "4/4 ✅"],
        ["EscalationAgent",                  "4",  "4/4 ✅"],
        ["FallbackAgent / Self-Healing",     "5",  "5/5 ✅"],
        ["AuditLoggerAgent",                 "2",  "2/2 ✅"],
        ["All 6 Hooks",                      "6",  "6/6 ✅"],
        ["Compliance Risk Detection",        "3",  "3/3 ✅"],
        ["End-to-End Pipeline (9 steps)",    "10", "10/10 ✅"],
    ]
)
tip(doc, "Say: 'Every component was tested independently and together. "
         "Everything passed first time.'")


# ── 9. LOAD TESTING RESULTS ───────────────────────────────────────────────────
section_title(doc, "9️⃣", "Load Testing Results  [~20 seconds]", NAVY)
body(doc,
    "We ran a real multi-threaded load test against the full agent pipeline — "
    "not just the search function, but all 7 agents, 3 skills, and 6 hooks together."
)
mini_table(doc,
    ["Metric", "Result", "What it means"],
    [
        ["Total requests",    "60 / 60",   "Not a single failure"],
        ["Concurrent workers","10",        "10 employees asking questions simultaneously"],
        ["Median latency",    "8.8 ms",    "Under 9ms for a typical query — practically instant"],
        ["p99 latency",       "30.3 ms",   "Even worst-case is well under human perception"],
        ["Throughput",        "693 req/s", "Can handle 693 queries per second if needed"],
        ["Answer found rate", "90%",       "9 in 10 questions matched a policy"],
        ["Errors",            "0",         "Zero failures under concurrent load"],
    ]
)
tip(doc, "Put it simply: 'Under a real concurrent load test, median response "
         "was under 9 milliseconds. That's faster than a blink.'")


# ── 10. DEPLOYMENT ARCHITECTURE ───────────────────────────────────────────────
section_title(doc, "🔟", "Deployment Architecture  [~20 seconds]", NAVY)
body(doc,
    "The system is flexible — it supports four deployment modes. "
    "Pick whichever suits the bank's IT setup:"
)
mini_table(doc,
    ["Mode", "How", "Best For"],
    [
        ["Local (Option A)",   "streamlit run app.py",                  "Single user, a personal workstation"],
        ["Docker (Option B)",  "docker run -p 8501:8501 bank-assistant","Isolated, consistent environment"],
        ["Network (Option C)", "Add --server.address=0.0.0.0",          "Whole team access on intranet"],
        ["MCP Server (Option D)","python mcp_server.py",                "AI agent or Claude Desktop integration"],
    ]
)
bullet(doc, [
    "Requirements: Python 3.10+, 512 MB RAM, 100 MB disk — any internal server",
    "No cloud account needed, no API keys, no internet",
    "All four modes share the same policy files, agents, and logs",
])
tip(doc, "'It runs anywhere inside the bank. No external dependencies. "
         "IT can deploy it in under 10 minutes.'")


# ── 11. SCREENSHOTS OF RESULTS ────────────────────────────────────────────────
section_title(doc, "📸", "Screenshots of Results  [~20 seconds]", NAVY)
body(doc,
    "We have 9 screenshots of the live application covering all core flows. "
    "Point to each one and name what's happening:"
)
mini_table(doc,
    ["Screenshot", "What to say"],
    [
        ["sc9_app_home.png",          "'This is the home screen — clean, professional, all 5 policies visible in the sidebar'"],
        ["sc1_loan_policy_search.png","'Employee types a loan question — answer comes back in under a second with the source badge'"],
        ["sc2_kyc_search.png",        "'KYC question — same experience, different policy matched automatically'"],
        ["sc3_multi_policy.png",      "'This one spans two policies — the system detected that and merged both answers'"],
        ["sc4_summary.png",           "'Ask for a summary — you get 5 clean bullet points, not a wall of text'"],
        ["sc5_escalation.png",        "'Employee needs to escalate — the system drafts the email for them automatically'"],
        ["sc6_compliance_alert.png",  "'Risk keyword detected — yellow alert fires immediately, logged for compliance team'"],
        ["sc7_fallback_selfheal.png", "'Nothing matched — self-healing fallback kicks in, contact details shown'"],
        ["sc8_account_opening.png",   "'Account opening policy — exact paragraph returned with full traceability'"],
    ]
)
tip(doc, "You don't need to show all 9. Pick 3–4 that tell the best story: "
         "home, loan search, multi-policy, compliance alert.")


# ── 12. BUSINESS IMPACT ───────────────────────────────────────────────────────
section_title(doc, "🎯", "Business Impact  [~20 seconds]", GREEN)
body(doc,
    "This is the slide that lands. Here's the before and after:"
)
mini_table(doc,
    ["Before", "After"],
    [
        ["5–10 minutes per policy lookup",           "Under 1 second — every time"],
        ["Different employees, different answers",   "Same approved answer, every time"],
        ["No audit trail",                           "Every query logged with timestamp, session, policy"],
        ["No compliance screening",                  "14 risk keywords auto-detected in real time"],
        ["Manual escalation drafting",               "Auto-generated escalation email, ready to send"],
        ["No AI integration possible",               "5 MCP tools — any AI agent can plug in now"],
        ["Cloud API dependency (typical AI tools)",  "100% offline — zero data leaves the bank"],
    ]
)

callout(doc,
    "If 100 employees each save 5 minutes a day → that's 500 minutes = over 8 hours "
    "of productivity recovered every single day.", GREEN)

body(doc,
    "Beyond time saved, the real win is consistency and confidence. "
    "Every employee, no matter how new, gives the same accurate, policy-backed answer. "
    "The compliance team gets an automatic audit trail without lifting a finger."
)
tip(doc, "Close with: 'We built something that makes every bank employee "
         "an expert on day one — and gives compliance a complete audit trail automatically.'")


# ── CLOSING NOTE ──────────────────────────────────────────────────────────────
doc.add_page_break()
section_title(doc, "💡", "Before You Present — Quick Reminders", ORANGE)
bullet(doc, [
    "You don't need to cover every detail — pick the highlights that your audience cares about",
    "If the audience is technical → spend more time on Sections 3, 4, 5, 9",
    "If the audience is business/management → spend more on Sections 1, 2, 12 (impact)",
    "If the audience is compliance/risk → lean into Sections 6, 7, 8",
    "Always end on Business Impact — it's the most memorable slide",
    "Screenshots are your best friend — show, don't tell",
    "If you get a hard question, it's okay to say 'great point, let me follow up' — "
    "you have all the detail in the full deliverable documents",
], ORANGE)

doc.add_paragraph()
final = doc.add_paragraph()
final.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = final.add_run(f"Prepared {TODAY}  ·  NexaBank Policy Assistant — Capstone Project")
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = LIGHT_GRAY


# ── SAVE ─────────────────────────────────────────────────────────────────────
out = "/home/labuser/ClaudeProject/CapstoneProject/deliverables/My_5Min_Presentation_Guide.docx"
doc.save(out)
print(f"✅  Saved: {out}")
